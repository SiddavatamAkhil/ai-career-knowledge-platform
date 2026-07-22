"""
Loads raw text out of user-supplied documents.

Supports: .txt, .md, .pdf, .docx
Each loaded document becomes a `Document` object carrying its source text
plus lightweight metadata (filename, page if applicable). Keeping metadata
attached from the very first step is what lets us cite sources later.
"""
from dataclasses import dataclass, field
from pathlib import Path
from typing import List


@dataclass
class Document:
    text: str
    metadata: dict = field(default_factory=dict)


def load_txt(path: Path) -> List[Document]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [Document(text=text, metadata={"source": path.name})]


def load_pdf(path: Path) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            docs.append(Document(text=page_text, metadata={"source": path.name, "page": i + 1}))
    return docs


def load_docx(path: Path) -> List[Document]:
    import docx

    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs)
    return [Document(text=text, metadata={"source": path.name})]


LOADERS = {
    ".txt": load_txt,
    ".md": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(path: str) -> List[Document]:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in LOADERS:
        raise ValueError(f"Unsupported file type: {ext}")
    return LOADERS[ext](p)


def load_directory(directory: str) -> List[Document]:
    """Load every supported file in a directory."""
    docs: List[Document] = []
    for p in sorted(Path(directory).glob("*")):
        if p.suffix.lower() in LOADERS:
            docs.extend(load_document(str(p)))
    return docs
